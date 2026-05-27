"""Export meeting minutes to Word (.docx) format."""

import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_styled_heading(doc: Document, text: str, level: int):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)


def _add_bullet(doc: Document, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def markdown_to_docx(analysis: str, meeting: dict) -> io.BytesIO:
    """Convert meeting analysis markdown to a Word document.

    Returns a BytesIO buffer containing the .docx file.
    """
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    title_text = meeting.get("title", "会议纪要")
    if title_text.startswith("_auto_"):
        title_text = "会议纪要"
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meeting metadata block
    meta_parts = []
    if meeting.get("meeting_time"):
        meta_parts.append(f"会议时间：{meeting['meeting_time']}")
    if meeting.get("location"):
        meta_parts.append(f"会议地点：{meeting['location']}")
    if meeting.get("participants"):
        meta_parts.append(f"参与人：{meeting['participants']}")
    if meeting.get("duration"):
        mins = int(meeting["duration"] // 60)
        secs = int(meeting["duration"] % 60)
        meta_parts.append(f"音频时长：{mins}分{secs}秒")

    if meta_parts:
        for part in meta_parts:
            p = doc.add_paragraph(part)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        doc.add_paragraph("")  # spacer

    # Parse markdown line by line
    lines = analysis.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # H1
        m = re.match(r"^#\s+(.+)$", stripped)
        if m and not stripped.startswith("## "):
            i += 1
            continue  # skip — already used as document title

        # H2
        m = re.match(r"^##\s+(.+)$", stripped)
        if m:
            _add_styled_heading(doc, m.group(1).strip(), level=1)
            i += 1
            continue

        # H3
        m = re.match(r"^###\s+(.+)$", stripped)
        if m:
            _add_styled_heading(doc, m.group(1).strip(), level=2)
            i += 1
            continue

        # Bullet list (- or *)
        m = re.match(r"^[-*]\s+(.+)$", stripped)
        if m:
            content = m.group(1)
            # Bold prefix like **负责人：**
            bm = re.match(r"\*\*(.+?)\*\*\s*(.*)", content)
            if bm:
                _add_bullet(doc, bm.group(2), bold_prefix=bm.group(1) + " ")
            else:
                _add_bullet(doc, _clean_inline_md(content))
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\d+\.\s+(.+)$", stripped)
        if m:
            content = m.group(1)
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, content)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$|^\*{3,}$", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run("─" * 40)
            run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            run.font.size = Pt(9)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _clean_inline_md(text: str) -> str:
    """Remove markdown inline formatting for plain text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _add_inline_runs(paragraph, text: str):
    """Parse inline markdown (bold, italic) and add styled runs."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        bm = re.match(r"\*\*(.+?)\*\*", part)
        if bm:
            run = paragraph.add_run(bm.group(1))
            run.bold = True
        else:
            paragraph.add_run(part)
